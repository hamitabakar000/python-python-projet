# -*- coding: utf-8 -*-
"""
Génère le rapport Word du projet Django PreSkool.

Prérequis : pip install python-docx pillow
Lancement (depuis ce dossier) : python generate_rapport.py

Pour utiliser VOS propres captures : remplacez les fichiers PNG dans ./captures/
(ou supprimez-les : le script recréera des illustrations par défaut).
"""

from __future__ import annotations

import os
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installez python-docx : pip install python-docx pillow")
    sys.exit(1)

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = None  # type: ignore

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CAPTURES_DIR = os.path.join(SCRIPT_DIR, "captures")
OUTPUT_NAME = "Rapport_Projet_Django_Hamit_Abakar_Moussa.docx"

# --- Métadonnées du rapport ---
AUTEUR = "Hamit Abakar Moussa"
ENCADRANT = "Mme Sara Ahsain"
GITHUB = "https://github.com/hamitabakar000/python-python-projet"
TITRE_PROJET = "Système de gestion scolaire PreSkool (Django)"
MODULE = "Développement web avancé Back-end (Python)"


def ensure_captures_dir() -> None:
    os.makedirs(CAPTURES_DIR, exist_ok=True)


def make_placeholder_image(path: str, title: str) -> None:
    """Crée une illustration type capture d'écran si Pillow est disponible."""
    if Image is None:
        return
    w, h = 1280, 720
    img = Image.new("RGB", (w, h), "#f1f5f9")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w, 72], fill="#4f46e5")
    try:
        font = ImageFont.truetype("arial.ttf", 28)
        font_small = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        font = ImageFont.load_default()
        font_small = font
    draw.text((32, 22), "PreSkool", fill="white", font=font)
    draw.rectangle([48, 100, w - 48, h - 48], outline="#cbd5e1", width=2, fill="white")
    draw.text((72, 130), title, fill="#0f172a", font=font)
    draw.text(
        (72, 200),
        "Illustration générée pour le rapport — remplacez ce fichier par une",
        fill="#64748b",
        font=font_small,
    )
    draw.text(
        (72, 230),
        "vraie capture d'écran de votre navigateur pour le rendu final.",
        fill="#64748b",
        font=font_small,
    )
    img.save(path, "PNG")


def ensure_default_captures() -> list[tuple[str, str]]:
    """Retourne la liste (fichier, légende). Crée des PNG si absents."""
    specs = [
        ("01_page_accueil.png", "Page d'accueil de l'application"),
        ("02_connexion.png", "Formulaire de connexion"),
        ("03_dashboard_admin.png", "Tableau de bord (vue administrateur)"),
        ("04_gestion_etudiants.png", "Exemple de gestion des étudiants"),
    ]
    ensure_captures_dir()
    result = []
    for fname, caption in specs:
        fpath = os.path.join(CAPTURES_DIR, fname)
        if not os.path.isfile(fpath) and Image is not None:
            make_placeholder_image(fpath, caption)
        result.append((fpath, caption))
    return result


def set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def add_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITRE_PROJET)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Rapport de projet").font.size = Pt(14)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Auteur : {AUTEUR}").font.size = Pt(12)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(f"Encadrante : {ENCADRANT}").font.size = Pt(12)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.add_run(f"Module : {MODULE}").font.size = Pt(11)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.add_run("Année universitaire 2025-2026").font.size = Pt(11)

    doc.add_page_break()


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover(doc)

    doc.add_heading("Remerciements", level=1)
    doc.add_paragraph(
        f"Je tiens à remercier {ENCADRANT} pour son encadrement, ses conseils "
        "et la qualité du cours dispensé dans le cadre de ce module. "
        "Ce projet m'a permis de mettre en pratique les concepts du framework Django "
        "et du développement web côté serveur."
    )

    doc.add_page_break()
    doc.add_heading("Table des matières", level=1)
    toc = [
        "1. Introduction",
        "2. Contexte et objectifs du projet",
        "3. Analyse des besoins fonctionnels",
        "4. Environnement et choix techniques",
        "5. Architecture de l'application (MVT)",
        "6. Modèles de données et modules Django",
        "7. Réalisation : fonctionnalités principales",
        "8. Captures d'écran",
        "9. Tests, déploiement local et données de démonstration",
        "10. Conclusion et perspectives",
        "11. Références et dépôt GitHub",
    ]
    for line in toc:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Ce document présente le travail réalisé dans le cadre d'un projet "
        "de développement d'une application web de gestion scolaire, nommée PreSkool, "
        "en utilisant le langage Python et le framework Django. "
        "L'objectif est de centraliser des informations pédagogiques et administratives "
        "(étudiants, enseignants, départements, matières, emploi du temps, examens, etc.) "
        "dans une interface web cohérente, avec une authentification et des rôles distincts."
    )

    doc.add_heading("2. Contexte et objectifs du projet", level=1)
    doc.add_paragraph(
        "Les établissements ont besoin d'outils pour suivre les inscriptions, "
        "le personnel, les cours et les résultats. Ce projet vise à :"
    )
    for item in [
        "Proposer une application modulaire, évolutive, basée sur une architecture MVT claire.",
        "Sécuriser l'accès par comptes utilisateurs (administrateur, enseignant, étudiant).",
        "Offrir des écrans de consultation et de gestion alignés sur un thème Bootstrap (template PreSkool).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Analyse des besoins fonctionnels", level=1)
    doc.add_paragraph("Les besoins couverts incluent notamment :")
    needs = [
        "Authentification : inscription, connexion, déconnexion, demande de réinitialisation de mot de passe.",
        "Gestion des rôles via un modèle utilisateur personnalisé (CustomUser).",
        "Gestion des étudiants et des parents (liaisons entre entités).",
        "Gestion des enseignants et des départements.",
        "Gestion des matières rattachées aux départements et aux enseignants.",
        "Gestion des jours fériés, de l'emploi du temps, des examens et des résultats.",
        "Tableaux de bord différenciés selon le profil (admin / enseignant).",
    ]
    for n in needs:
        doc.add_paragraph(n, style="List Bullet")

    doc.add_heading("4. Environnement et choix techniques", level=1)
    doc.add_paragraph(
        "L'application est développée en Python avec Django 6, qui fournit l'ORM, "
        "le système d'administration, le routage URL, les vues et le moteur de templates. "
        "La base de données par défaut est SQLite, adaptée au prototypage et au rendu académique. "
        "Les interfaces s'appuient sur des gabarits HTML/CSS et Bootstrap (thème PreSkool). "
        "Le traitement d'images pour les photos (enseignants, etc.) utilise Pillow."
    )

    doc.add_heading("5. Architecture de l'application (MVT)", level=1)
    doc.add_paragraph(
        "Django suit le pattern Model-View-Template : les modèles décrivent les tables "
        "et les relations ; les vues (fonctions ou classes) traitent la requête HTTP et "
        "préparent le contexte ; les templates génèrent le HTML. "
        "Les URL principales regroupent les applications `home_auth`, `faculty`, `student`, "
        "`teacher`, `department`, `subject`, `holiday`, `timetable` et `exam`."
    )

    doc.add_heading("6. Modèles de données et modules Django", level=1)
    doc.add_paragraph(
        "Le cœur métier repose sur des entités reliées par des clés étrangères : "
        "Department, Teacher, Student, Parent, Subject, Holiday, Timetable, Exam, ExamResult, etc. "
        "Le modèle CustomUser étend AbstractUser avec des indicateurs de rôle "
        "(is_admin, is_teacher, is_student) et des champs adaptés au projet."
    )

    doc.add_heading("7. Réalisation : fonctionnalités principales", level=1)
    doc.add_paragraph(
        "Les parcours utilisateur incluent la connexion avec redirection vers les tableaux de bord, "
        "la navigation dans les sections étudiants / enseignants / matières, "
        "et l'administration Django pour les opérations de maintenance. "
        "Un script seed.py permet de charger des données de démonstration (comptes de test, "
        "départements, sujets fictifs, etc.)."
    )

    doc.add_heading("8. Captures d'écran", level=1)
    doc.add_paragraph(
        "Les figures ci-dessous illustrent l'interface de l'application. "
        "Pour le rendu final, vous pouvez remplacer les fichiers du dossier `rapport/captures` "
        "par de véritables captures de votre navigateur (mêmes noms de fichiers), puis relancer ce script."
    )

    captures = ensure_default_captures()
    for idx, (fpath, caption) in enumerate(captures, start=1):
        doc.add_paragraph()
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap_p.add_run(f"Figure {idx} — {caption}")
        r.italic = True
        r.font.size = Pt(10)
        if os.path.isfile(fpath):
            try:
                doc.add_picture(fpath, width=Cm(15))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph(f"(Impossible d'inclure l'image : {fpath})")
        else:
            doc.add_paragraph(f"(Fichier manquant : {fpath})")

    doc.add_heading("9. Tests, déploiement local et données de démonstration", level=1)
    doc.add_paragraph(
        "En local, on exécute `python manage.py migrate` puis `python manage.py runserver` "
        "depuis le dossier `school`. Le script `python seed.py` permet de recréer des comptes de test, "
        "notamment administrateur (admin@preskool.com), enseignant et étudiant, "
        "conformément au README du projet."
    )

    doc.add_heading("10. Conclusion et perspectives", level=1)
    doc.add_paragraph(
        "Ce projet démontre la mise en œuvre d'une application Django multi-applications "
        "avec authentification personnalisée et domaine métier riche. "
        "Des prolongements possibles incluent : déploiement sur un serveur (PostgreSQL, HTTPS), "
        "envoi réel des e-mails de réinitialisation, tests automatisés et API REST."
    )

    doc.add_heading("11. Références et dépôt GitHub", level=1)
    p = doc.add_paragraph("Code source et historique du projet : ")
    add_hyperlink(p, GITHUB, GITHUB)
    doc.add_paragraph()
    doc.add_paragraph(f"Dépôt : {GITHUB}")

    # Pied de page métadonnées (section propiétaire)
    core_props = doc.core_properties
    core_props.author = AUTEUR
    core_props.title = TITRE_PROJET

    return doc


def main() -> None:
    doc = build_document()
    out_path = os.path.join(SCRIPT_DIR, OUTPUT_NAME)
    doc.save(out_path)
    print(f"Rapport généré : {out_path}")


if __name__ == "__main__":
    main()
